"""RAMS (Risk Assessment Method Statement) generator.

Reads site_data.csv and produces RAMS.docx — a formatted Word document suitable
for HSE submission. Requires python-docx.

Usage:
    python scripts/hse_rams.py [--csv data/site_data.csv] [--out RAMS.docx]
    python scripts/hse_rams.py --project "Grid B5 Drainage Works" --site "Unit 7, Industrial Estate"

CSV columns (see data/site_data.csv for example):
    activity, hazard, persons_at_risk, severity, likelihood, controls,
    residual_severity, residual_likelihood, ppe, responsible
"""
from __future__ import annotations

import argparse
import csv
import datetime
import sys
from pathlib import Path
from typing import Any, Dict, List


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CSV = ROOT / "data" / "site_data.csv"
DEFAULT_OUT = ROOT / "RAMS.docx"

RISK_LABEL = {
    1: "Very Low", 2: "Low", 3: "Medium", 4: "High", 5: "Very High"
}

# RGB colours used for risk-level cells
RISK_COLOURS: dict[int, str] = {
    1: "00B050",  # green
    2: "92D050",  # light green
    3: "FFFF00",  # yellow
    4: "FFC000",  # amber
    5: "FF0000",  # red
}


def _risk_score(severity: int, likelihood: int) -> int:
    return severity * likelihood


def _risk_colour(score: int) -> str:
    if score <= 4:
        return RISK_COLOURS[1]
    if score <= 8:
        return RISK_COLOURS[2]
    if score <= 12:
        return RISK_COLOURS[3]
    if score <= 16:
        return RISK_COLOURS[4]
    return RISK_COLOURS[5]


def load_activities(csv_path: Path) -> List[Dict[str, Any]]:
    rows = []
    with csv_path.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append({
                "activity": row.get("activity", "").strip(),
                "hazard": row.get("hazard", "").strip(),
                "persons_at_risk": row.get("persons_at_risk", "").strip(),
                "severity": int(row.get("severity", 3)),
                "likelihood": int(row.get("likelihood", 3)),
                "controls": row.get("controls", "").strip(),
                "residual_severity": int(row.get("residual_severity", 2)),
                "residual_likelihood": int(row.get("residual_likelihood", 1)),
                "ppe": row.get("ppe", "").strip(),
                "responsible": row.get("responsible", "").strip(),
            })
    return rows


def generate_rams(
    csv_path: Path = DEFAULT_CSV,
    out_path: Path = DEFAULT_OUT,
    project_name: str = "Site Works",
    site_address: str = "TBC",
    prepared_by: str = "Site Manager",
    review_date: str | None = None,
) -> Path:
    """Generate RAMS.docx from csv_path. Returns the output path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    if not csv_path.exists():
        raise FileNotFoundError(f"Site data CSV not found: {csv_path}")

    activities = load_activities(csv_path)
    if not activities:
        raise ValueError("No rows found in site data CSV")

    if review_date is None:
        review_date = (datetime.date.today() + datetime.timedelta(days=90)).strftime("%d/%m/%Y")

    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── title block ──
    title = doc.add_heading("RISK ASSESSMENT & METHOD STATEMENT (RAMS)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_table = doc.add_table(rows=4, cols=4)
    meta_table.style = "Table Grid"
    labels = [
        ("Project:", project_name, "Site Address:", site_address),
        ("Prepared by:", prepared_by, "Date:", datetime.date.today().strftime("%d/%m/%Y")),
        ("Review date:", review_date, "Document ref:", f"RAMS-{datetime.date.today().strftime('%Y%m%d')}"),
        ("Issue:", "1.0", "Status:", "APPROVED FOR USE"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(labels):
        row = meta_table.rows[i]
        for j, txt in enumerate((l1, v1, l2, v2)):
            cell = row.cells[j]
            run = cell.paragraphs[0].add_run(txt)
            if j in (0, 2):
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ── scope paragraph ──
    scope = doc.add_paragraph()
    scope.add_run("Scope of Works: ").bold = True
    scope.add_run(
        f"This RAMS covers the activities listed below for {project_name}. "
        "All operatives must be briefed on this document before commencing work. "
        "The site supervisor must ensure controls are in place and record signatures overleaf."
    )
    scope.runs[-1].font.size = Pt(9)

    doc.add_paragraph()

    # ── risk matrix header ──
    doc.add_heading("Risk Matrix (Severity × Likelihood)", level=2)

    matrix_table = doc.add_table(rows=2, cols=6)
    matrix_table.style = "Table Grid"
    headers = ["Score", "1–4\nVery Low", "5–8\nLow", "9–12\nMedium", "13–16\nHigh", "17–25\nVery High"]
    colours = [None, "00B050", "92D050", "FFFF00", "FFC000", "FF0000"]
    for j, (hdr, colour) in enumerate(zip(headers, colours)):
        c = matrix_table.rows[0].cells[j]
        c.paragraphs[0].add_run(hdr).bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        if colour:
            _shade_cell(c, colour)
    for j in range(6):
        c = matrix_table.rows[1].cells[j]
        c.paragraphs[0].add_run("SxL").font.size = Pt(8)

    doc.add_paragraph()

    # ── main RAMS table ──
    doc.add_heading("Activity Risk Assessments", level=2)

    COLS = [
        "Activity", "Hazard", "Persons\nAt Risk",
        "Sev", "Like", "Risk\nScore", "Risk\nRating",
        "Control Measures",
        "Res.\nSev", "Res.\nLike", "Res.\nScore", "Res.\nRating",
        "PPE Required", "Responsible",
    ]
    rams_table = doc.add_table(rows=1 + len(activities), cols=len(COLS))
    rams_table.style = "Table Grid"

    # header row
    hdr_row = rams_table.rows[0]
    for j, h in enumerate(COLS):
        cell = hdr_row.cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(cell, "D9D9D9")

    # data rows
    for i, act in enumerate(activities):
        row = rams_table.rows[i + 1]
        score = _risk_score(act["severity"], act["likelihood"])
        res_score = _risk_score(act["residual_severity"], act["residual_likelihood"])

        values = [
            act["activity"],
            act["hazard"],
            act["persons_at_risk"],
            str(act["severity"]),
            str(act["likelihood"]),
            str(score),
            _risk_label_from_score(score),
            act["controls"],
            str(act["residual_severity"]),
            str(act["residual_likelihood"]),
            str(res_score),
            _risk_label_from_score(res_score),
            act["ppe"],
            act["responsible"],
        ]
        risk_col = 6
        res_risk_col = 11

        for j, val in enumerate(values):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(7)
            if j in (3, 4, 5, 6, 8, 9, 10, 11):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == risk_col:
                _shade_cell(cell, _risk_colour(score))
            if j == res_risk_col:
                _shade_cell(cell, _risk_colour(res_score))

    doc.add_paragraph()

    # ── sign-off section ──
    doc.add_heading("Operative Briefing & Sign-Off", level=2)
    signoff_para = doc.add_paragraph(
        "I confirm I have read, understood, and will comply with this RAMS. "
        "Any concerns must be raised with the site supervisor before work begins."
    )
    signoff_para.runs[0].font.size = Pt(9)

    signoff_table = doc.add_table(rows=11, cols=4)
    signoff_table.style = "Table Grid"
    sh = signoff_table.rows[0]
    for j, hdr in enumerate(("Name (print)", "Signature", "Role", "Date")):
        run = sh.cells[j].paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
    for row in signoff_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].add_run(" ")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _risk_label_from_score(score: int) -> str:
    if score <= 4:
        return "Very Low"
    if score <= 8:
        return "Low"
    if score <= 12:
        return "Medium"
    if score <= 16:
        return "High"
    return "Very High"


def _shade_cell(cell: Any, hex_colour: str) -> None:
    """Apply a background fill to a table cell via direct XML manipulation."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Generate RAMS.docx from site_data.csv")
    parser.add_argument("--csv", type=Path, default=DEFAULT_CSV, help="Path to site_data.csv")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT, help="Output .docx path")
    parser.add_argument("--project", default="Site Works", help="Project name")
    parser.add_argument("--site", default="TBC", help="Site address")
    parser.add_argument("--prepared-by", default="Site Manager", help="Prepared by")
    args = parser.parse_args(argv)

    out = generate_rams(
        csv_path=args.csv,
        out_path=args.out,
        project_name=args.project,
        site_address=args.site,
        prepared_by=args.prepared_by,
    )
    print(f"RAMS written to: {out}")


if __name__ == "__main__":
    main()
